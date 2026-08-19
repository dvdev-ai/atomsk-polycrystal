#!/usr/bin/env python3
"""Текстовый отчёт для устной защиты (Word)."""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Cm

OUT = Path.home() / "Downloads" / "ATOMSK_справка_для_защиты.docx"


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        r.font.name = "Times New Roman"


def add_para(doc: Document, text: str, bold: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    run.bold = bold
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15


def add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text, style="List Bullet")
    for r in p.runs:
        r.font.name = "Times New Roman"
        r.font.size = Pt(12)


def main() -> None:
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Cm(2)
    sec.bottom_margin = Cm(2)
    sec.left_margin = Cm(2.5)
    sec.right_margin = Cm(1.5)

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run(
        "Справка для устной защиты курсовой\n"
        "ATOMSK · LAMMPS · OVITO · диаграммы σ–ε"
    )
    r.bold = True
    r.font.size = Pt(14)
    r.font.name = "Times New Roman"

    add_para(
        doc,
        "Документ поясняет, за что отвечает каждый рисунок в презентации и в папке figures/. "
        "Численные расчёты и графики σ–ε не пересчитывались — описание соответствует "
        "готовым результатам проекта atomsk-polycrystal-tension.",
    )

    add_heading(doc, "1. Общая логика работы", 1)
    add_para(
        doc,
        "Цепочка «построение модели → молекулярная динамика → визуализация → графики» "
        "стандартна в материаловедении:",
    )
    add_bullet(doc, "ATOMSK — создаёт поликристалл (8 зёрен, куб 60×60×60 Å) из seed-ячейки и файла параметров Вороного.")
    add_bullet(doc, "LAMMPS — минимизирует энергию, релаксирует образец при заданной T, тянет по оси X, пишет stress_strain.dat.")
    add_bullet(doc, "OVITO — показывает зёрна (свойство grainID) и деформированное состояние по траектории dump.")
    add_bullet(doc, "Python (plot_stress_strain.py) — строит диаграммы напряжение–деформация для отчёта и слайдов.")

    add_heading(doc, "2. Графики σ–ε (не менять)", 1)
    add_para(
        doc,
        "Графики в figures/ — главный количественный результат. Их смысл для защиты:",
    )

    add_heading(doc, "2.1. stress_strain_Cu.png и stress_strain_Al.png", 2)
    add_bullet(doc, "Ось X — инженерная деформация ε в процентах (одноосное растяжение по X).")
    add_bullet(doc, "Ось Y — осевое напряжение σ_xx в ГПа из выходного файла LAMMPS.")
    add_bullet(doc, "Три кривые на каждом рисунке — температуры 0 K (в MD ≈ 1 K), 300 K и 600 K.")
    add_bullet(doc, "Физический смысл: с ростом T максимальное напряжение обычно снижается — материал «мягче».")
    add_bullet(doc, "На защите: «Мы получили зависимость σ(ε) для поликристаллической меди/алюминия при трёх температурах».")

    add_heading(doc, "2.2. stress_strain_combined.png", 2)
    add_bullet(doc, "Два подграфика Cu и Al на одном слайде — для сравнения материалов при одинаковых условиях MD.")
    add_bullet(doc, "Удобно на слайде «Результаты», когда мало времени.")

    add_heading(doc, "2.3. summary_table.png", 2)
    add_bullet(doc, "Таблица σ_max и ε_max по каждому расчёту — краткая численная сводка без чтения всех .dat файлов.")
    add_bullet(doc, "Используйте как опору, если спросят «какие порядки величин получились».")

    add_heading(doc, "3. Скриншоты OVITO («кубы»)", 1)

    add_heading(doc, "3.1. Первый и второй кадр — всё в порядке", 2)
    add_bullet(doc, "ovito_Cu_polycrystal.png — начальная модель меди: 8 зёрен, раскраска по grainID.")
    add_bullet(doc, "ovito_Al_polycrystal.png — то же для алюминия.")
    add_bullet(doc, "На слайде говорите: «Видна зернистая структура, границы зёрен, метод Вороного в ATOMSK».")
    add_bullet(doc, "Белый фон и перспективная камера — «образец в рабочем объёме симуляции».")

    add_heading(doc, "3.2. Третий кадр — что не так и как исправить", 2)
    add_para(
        doc,
        "Слайд «Деформированное состояние» часто выглядит грубо, если:",
        bold=True,
    )
    add_bullet(doc, "в кадр попадают два периодических изображения ящика (визуально — «два куба»);")
    add_bullet(doc, "раскраска по типу атома (Cu/Al) вместо grainID — стиль не совпадает с первыми слайдами;")
    add_bullet(doc, "берётся кадр с сильной неоднородной деформацией или серый фон — контраст с Cu/Al.")
    add_para(
        doc,
        "Исправленный кадр: ovito_tension_Cu_T300K.png — один объём, белый фон, zoom_all, "
        "без файлов _top и _mid. В траектории LAMMPS нет grainID, поэтому раскраска по "
        "координате X (направление растяжения) — на слайде так и говорите.",
    )

    add_heading(doc, "4. Структура презентации (что говорить)", 1)
    slides = [
        ("Титул", "Тема, ФИО, кафедра."),
        ("Цель", "Изучить ATOMSK, построить поликристаллы, получить σ–ε, визуализировать."),
        ("ATOMSK", "Командная строка, polycrystal, экспорт в LAMMPS."),
        ("Методика", "60 Å, 8 зёрен, EAM, NPT, fix deform."),
        ("OVITO Cu / Al", "Два «куба» — зернистость, grainID."),
        ("OVITO деформация", "Один кадр: как меняется структура при растяжении."),
        ("Графики σ–ε", "Количественный результат; влияние температуры."),
        ("Выводы", "Пайплайн работает; для эксперимента нужны больший образец и меньшая ε̇."),
    ]
    for title, note in slides:
        add_para(doc, f"• {title} — {note}")

    add_heading(doc, "5. Типичные вопросы на защите", 1)
    qa = [
        ("Почему ε̇ = 10⁹ s⁻¹?", "В MD нужны наносекунды; это стандартный компромисс, не лабораторная скорость испытания."),
        ("Почему 0 K?", "Классическая MD: используем 1 K после минимизации как «нулевую» точку."),
        ("Зачем поликристалл?", "Реалистичнее однокристалла; границы зёрен влияют на σ–ε."),
        ("Чем ATOMSK лучше ручной сборки?", "Быстро, воспроизводимо, сразу grainID для OVITO."),
    ]
    for q, a in qa:
        add_para(doc, q, bold=True)
        add_para(doc, a)

    doc.save(OUT)
    print(f"Сохранено: {OUT}")


if __name__ == "__main__":
    main()
